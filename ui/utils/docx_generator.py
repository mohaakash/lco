import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

class DocxReportGenerator:
    def __init__(self, output_path):
        self.output_path = output_path
        self.doc = Document()
        self._setup_styles()

    def _setup_styles(self):
        # Define custom colors
        self.color_primary = RGBColor(0, 139, 139)  # Dark Cyan #008B8B
        self.color_secondary = RGBColor(0, 59, 59)  # Darker Cyan #003b3b
        self.color_text = RGBColor(34, 34, 34)      # #222222
        self.color_subtle = RGBColor(102, 102, 102) # #666666

        # Update Normal style
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Helvetica Neue'
        font.size = Pt(11)
        font.color.rgb = self.color_text

        # Heading 1
        h1 = self.doc.styles['Heading 1']
        h1.font.name = 'Helvetica Neue'
        h1.font.size = Pt(24)
        h1.font.color.rgb = self.color_secondary
        h1.paragraph_format.space_before = Pt(24)
        h1.paragraph_format.space_after = Pt(12)

        # Heading 2
        h2 = self.doc.styles['Heading 2']
        h2.font.name = 'Helvetica Neue'
        h2.font.size = Pt(18)
        h2.font.color.rgb = self.color_secondary
        h2.paragraph_format.space_before = Pt(20)
        h2.paragraph_format.space_after = Pt(10)

        # Heading 3
        h3 = self.doc.styles['Heading 3']
        h3.font.name = 'Helvetica Neue'
        h3.font.size = Pt(14)
        h3.font.color.rgb = self.color_primary
        h3.paragraph_format.space_before = Pt(14)
        h3.paragraph_format.space_after = Pt(8)
        
        # Title Style
        if 'Title' in self.doc.styles:
            title = self.doc.styles['Title']
            title.font.name = 'Helvetica Neue'
            title.font.size = Pt(36)
            title.font.color.rgb = self.color_primary
            title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title.paragraph_format.space_after = Pt(12)

        # Subtitle Style
        if 'Subtitle' not in self.doc.styles:
            subtitle = self.doc.styles.add_style('Subtitle', WD_STYLE_TYPE.PARAGRAPH)
        else:
            subtitle = self.doc.styles['Subtitle']
        subtitle.font.name = 'Helvetica Neue'
        subtitle.font.size = Pt(16)
        subtitle.font.color.rgb = self.color_subtle
        subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.paragraph_format.space_after = Pt(24)

    def generate(self, data):
        self._create_cover_page(data)
        self.doc.add_page_break()
        
        self._create_personal_details(data)
        self._create_toc()
        self.doc.add_page_break()
        
        self._create_elemental_analysis(data)
        self.doc.add_page_break()
        
        self._create_daily_plan(data)
        self.doc.add_page_break()
        
        self._create_quality_assessment(data)
        self.doc.add_page_break()
        
        self._create_summary(data)

        self.doc.save(self.output_path)

    def _create_cover_page(self, data):
        # Logo
        logo_path = os.path.join('images', 'logo.png')
        if os.path.exists(logo_path):
            # Center the logo
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run()
            r.add_picture(logo_path, width=Inches(2))

        # Title
        self.doc.add_paragraph("Elemental Balance Assessment", style='Title')
        self.doc.add_paragraph("Personalised wellness guidance", style='Subtitle')

        # Cover Image
        cover_path = os.path.join('images', 'coverimage.png')
        if os.path.exists(cover_path):
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run()
            r.add_picture(cover_path, width=Inches(6))

        # Meta info
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(24)
        
        date_str = datetime.now().strftime("%B %d, %Y")
        p.add_run(date_str).font.color.rgb = self.color_subtle
        
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("Generated for: ").font.color.rgb = self.color_subtle
        
        personal = data.get('Personal', {})
        name = personal.get('name', 'Valued Client')
        run = p.add_run(name)
        run.bold = True
        run.font.color.rgb = self.color_text

    def _create_personal_details(self, data):
        self.doc.add_heading('1. Personal Details', level=1)
        
        personal = data.get('Personal', {})
        p = self.doc.add_paragraph()
        
        def add_field(label, value):
            run = p.add_run(f"{label}: ")
            run.font.color.rgb = self.color_subtle
            run = p.add_run(f"{value}\n")
            run.bold = True
        
        add_field("Name", personal.get('name', ''))
        add_field("Date of Birth", personal.get('date_of_birth', ''))
        add_field("Time of Birth", personal.get('time_of_birth', ''))
        add_field("Place of Birth", personal.get('place_of_birth', ''))
        add_field("Phone", personal.get('phone', ''))
        add_field("Email", personal.get('email', ''))

    def _create_toc(self):
        self.doc.add_heading('2. Table of Contents', level=1)
        # Static TOC for now as python-docx doesn't support updating dynamic TOC fields easily
        toc_items = [
            "1. Personal Details",
            "2. Elemental Analysis",
            "3. Daily Plan",
            "4. Elemental Quality Assessment"
        ]
        for item in toc_items:
            p = self.doc.add_paragraph(item)
            p.paragraph_format.left_indent = Inches(0.5)

    def _create_elemental_analysis(self, data):
        self.doc.add_heading('3. Elemental Analysis', level=1)
        
        # Element Percentages
        ep = data.get('Element_Percentages', {})
        if ep:
            self.doc.add_heading('Element Percentages', level=3)
            p = self.doc.add_paragraph()
            for k in ['Fire', 'Earth', 'Air', 'Water']:
                if k in ep:
                    run = p.add_run(f"{k}: ")
                    run.bold = True
                    p.add_run(f"{ep[k]}%   ")

        # Elements
        # Prefer Element_Descriptions (new structure) -> Elemental_Analysis (mapped)
        ea = data.get('Element_Descriptions') or data.get('Elemental_Analysis') or {}
        
        for el_name, el_data in ea.items():
            # Extract data based on structure
            title = ''
            content_raw = ''
            status = ''
            percentage = ''
            
            if isinstance(el_data, dict):
                title = el_data.get('Title') or el_data.get('Classification') or ''
                content_raw = el_data.get('Content') or el_data.get('Description') or ''
                status = el_data.get('Status', '')
                percentage = el_data.get('Percentage', '')
            else:
                content_raw = str(el_data)

            # Header
            header_text = f"{el_name}"
            if percentage:
                header_text += f" - {percentage}%"
            if status:
                header_text += f" ({status})"
            
            self.doc.add_heading(header_text, level=2)
            
            if title and title != el_name:
                self.doc.add_heading(title, level=3)

            # Process Content
            description = ""
            scientific = ""
            imbalance = ""
            remedies = {}

            if isinstance(content_raw, dict):
                # Nested JSON structure
                description = content_raw.get('The Fire Element') or content_raw.get('The Earth Element') or \
                              content_raw.get('The Air Element') or content_raw.get('The Water Element') or \
                              content_raw.get('Description') or ""
                scientific = content_raw.get('Scientific Correlation') or ""
                
                imbalance_parts = []
                if content_raw.get("What Low Fire Feels Like—and How It Holds You Back"):
                    imbalance_parts.append(content_raw.get("What Low Fire Feels Like—and How It Holds You Back"))
                if content_raw.get("What Excess Fire Feels Like—and Why You Need to Rein It In"):
                    imbalance_parts.append(content_raw.get("What Excess Fire Feels Like—and Why You Need to Rein It In"))
                if content_raw.get("Imbalance Effects"):
                    imbalance_parts.append(content_raw.get("Imbalance Effects"))
                imbalance = "\n".join(imbalance_parts)

                remedies = {
                    "Diet": content_raw.get("Diet", ""),
                    "Lifestyle & Exercise": content_raw.get("Lifestyle and Exercise", ""),
                    "Herbal / Energy Support": content_raw.get("Gems, Flower Remedies, and Aromas") or \
                                                content_raw.get("Herbs") or \
                                                content_raw.get("Crystals, Gems, and Herbal Remedies") or ""
                }
            elif isinstance(content_raw, str):
                description = content_raw
                # Try legacy fields
                if isinstance(el_data, dict):
                    scientific = el_data.get('Scientific_Correlation', '')
                    imbalance = el_data.get('Imbalance_Effects', '')
                    remedies = el_data.get('Remedies', {})
                    # Rename keys for display if needed
                    if 'Lifestyle_and_Exercise' in remedies:
                        remedies['Lifestyle & Exercise'] = remedies.pop('Lifestyle_and_Exercise')
                    if 'Herbal_or_Energy_Support' in remedies:
                        remedies['Herbal / Energy Support'] = remedies.pop('Herbal_or_Energy_Support')

            if description:
                self.doc.add_heading('Description', level=3)
                self.doc.add_paragraph(description)
            
            if scientific:
                self.doc.add_heading('Scientific Correlation', level=3)
                self.doc.add_paragraph(scientific)
                
            if imbalance:
                self.doc.add_heading('Imbalance Effects', level=3)
                self.doc.add_paragraph(imbalance)
                
            if remedies and any(remedies.values()):
                self.doc.add_heading('Remedies', level=3)
                table = self.doc.add_table(rows=0, cols=2)
                table.style = 'Table Grid'
                for k, v in remedies.items():
                    if v:
                        row = table.add_row()
                        row.cells[0].text = k
                        row.cells[1].text = str(v)
                        # Bold the key
                        row.cells[0].paragraphs[0].runs[0].bold = True

    def _create_daily_plan(self, data):
        self.doc.add_heading('4. Daily Plan', level=1)
        
        dg = data.get('Daily_Routine') or data.get('Daily_Guideline') or {}
        
        for period, content in dg.items():
            self.doc.add_heading(period, level=2)
            
            if isinstance(content, dict):
                table = self.doc.add_table(rows=0, cols=2)
                table.style = 'Table Grid'
                for k, v in content.items():
                    if v:
                        row = table.add_row()
                        # Clean key
                        key_clean = k.replace('_', ' ')
                        row.cells[0].text = key_clean
                        row.cells[1].text = str(v)
                        row.cells[0].paragraphs[0].runs[0].bold = True
            else:
                self.doc.add_paragraph(str(content))

    def _create_quality_assessment(self, data):
        self.doc.add_heading('5. Elemental Quality Assessment', level=1)
        
        # Fixed Intro Text
        intro_text = (
            "The distribution on planets in the zodiacal wheel shows how we respond to stimuli, "
            "and especially how we act under tension. The presence of planets in a particular sign "
            "indicates how that sign expresses their element, meaning the mode of behavior they use "
            "to express the element needs. There are three qualities: Cardinal, Fixed, and Mutable."
        )
        self.doc.add_paragraph(intro_text)
        
        # Modalities Percentages
        mp = data.get('Modalities_Percentages', {})
        if mp:
            self.doc.add_heading('Modalities Percentages', level=3)
            p = self.doc.add_paragraph()
            for k in ['Cardinal', 'Fixed', 'Mutable']:
                if k in mp:
                    run = p.add_run(f"{k}: ")
                    run.bold = True
                    p.add_run(f"{mp[k]}%   ")

        # Modalities
        mods = data.get('Modality_Descriptions') or data.get('Modalities') or {}
        
        for mname, mcontent in mods.items():
            pct = ''
            title = ''
            content = ''
            
            if isinstance(mcontent, dict):
                pct = mcontent.get('Percentage', '')
                title = mcontent.get('Title', '')
                
                inner_content = mcontent.get('Content')
                if isinstance(inner_content, dict):
                     # Render key-value pairs
                    lines = []
                    for k, v in inner_content.items():
                        lines.append(f"{k}: {v}")
                    content = "\n\n".join(lines)
                elif isinstance(inner_content, str):
                    content = inner_content
                else:
                    # Fallback
                    lines = []
                    for k, v in mcontent.items():
                        if k not in ('Percentage', 'Title', 'Content'):
                            lines.append(f"{k}: {v}")
                    content = "\n\n".join(lines) if lines else mcontent.get('Content', '')
            else:
                content = str(mcontent)

            header_text = f"{mname}"
            if pct:
                header_text += f" - {pct}%"
            
            self.doc.add_heading(header_text, level=2)
            
            if title:
                p = self.doc.add_paragraph(title)
                p.runs[0].bold = True
                p.runs[0].font.color.rgb = self.color_primary
            
            if content:
                # If content has newlines (from our join), split and add paragraphs
                for line in content.split('\n\n'):
                    if line.strip():
                        self.doc.add_paragraph(line.strip())

    def _create_summary(self, data):
        summary = data.get('Summary', '')
        if summary:
            self.doc.add_heading('Summary', level=2)
            self.doc.add_paragraph(summary)
            
        disclaimer = data.get('Disclaimer', '')
        if disclaimer:
            self.doc.add_heading('Disclaimer', level=2)
            p = self.doc.add_paragraph(disclaimer)
            p.style.font.size = Pt(9)
            p.style.font.color.rgb = self.color_subtle
