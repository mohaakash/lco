import docx
import sys

def read_docx(file_path):
    try:
        doc = docx.Document(file_path)
        print(f"--- Content of {file_path} ---")
        for para in doc.paragraphs:
            if para.text.strip():
                print(f"[{para.style.name}] {para.text}")
        
        print("\n--- Tables ---")
        for table in doc.tables:
            for row in table.rows:
                print(" | ".join([cell.text.strip() for cell in row.cells]))
            print("-" * 20)
            
    except Exception as e:
        print(f"Error reading docx: {e}")

if __name__ == "__main__":
    read_docx("docs/output.docx")
